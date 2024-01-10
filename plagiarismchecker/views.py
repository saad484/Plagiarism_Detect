from django.shortcuts import render
from django.http import HttpResponse
from plagiarismchecker.algorithm import main
from docx import *
from plagiarismchecker.algorithm import fileSimilarity
from PyPDF2 import PdfReader

# Create your views here.
#home
def home(request):
    return render(request, 'pc/index.html') 

#web search(Text)
def test(request):
    print("request is welcome test")
    print(request.POST['q'])  
    
    if request.POST['q']: 
        percent,link = main.findSimilarity(request.POST['q'])
        percent = round(percent,2)
    print("Output.....................!!!!!!!!",percent,link)
    return render(request, 'pc/index.html',{'link': link, 'percent': percent})

#web search file(.txt, .docx)
''' def filetest(request):
    value = ''    
    print(request.FILES['docfile'])
    if str(request.FILES['docfile']).endswith(".txt"):
        value = str(request.FILES['docfile'].read())

    elif str(request.FILES['docfile']).endswith(".docx"):
        document = Document(request.FILES['docfile'])
        for para in document.paragraphs:
            value += para.text

    elif str(request.FILES['docfile']).endswith(".pdf"):
        # creating a pdf file object 
        pdfFileObj = open(request.FILES['docfile'], 'rb') 

        # creating a pdf reader object 
        pdfReader = PyPDF2.PdfFileReader(pdfFileObj) 

        # printing number of pages in pdf file 
        print(pdfReader.numPages) 

        # creating a page object 
        pageObj = pdfReader.getPage(0) 

        # extracting text from page 
        print(pageObj.extractText()) 

        # closing the pdf file object 
        pdfFileObj.close() 


    percent,link = main.findSimilarity(value)
    print("Output...................!!!!!!!!",percent,link)
    return render(request, 'pc/index.html',{'link': link, 'percent': percent})
'''
def filetest(request):
    value = ''

    uploaded_file = request.FILES.get('docfile')

    if not uploaded_file:
        return HttpResponse("No file uploaded", status=400)

    file_name = uploaded_file.name.lower()

    if file_name.endswith(".txt"):
        value = uploaded_file.read().decode('utf-8')
    elif file_name.endswith(".docx"):
        document = Document(uploaded_file)
        value = "\n".join([para.text for para in document.paragraphs])
    elif file_name.endswith(".pdf"):
        try:
            pdf_reader = PdfReader(uploaded_file)
            value = ""
            for page_num in range(min(3, len(pdf_reader.pages))):  # Limit to first 3 pages for demo
                page = pdf_reader.pages[page_num]
                value += page.extract_text()
        except Exception as e:
            return HttpResponse(f"Error reading PDF: {e}", status=500)

    percent, link = main.findSimilarity(value)

    return render(request, 'pc/index.html', {'link': link, 'percent': round(percent, 2)})   

#text compare
def fileCompare(request):
    return render(request, 'pc/doc_compare.html') 

#two text compare(Text)
def twofiletest1(request):
    print("Submiited text for 1st and 2nd")
    print(request.POST['q1'])
    print(request.POST['q2'])

    if request.POST['q1'] != '' and request.POST['q2'] != '': 
        print("Got both the texts")
        result = fileSimilarity.findFileSimilarity(request.POST['q1'],request.POST['q2'])
    result = round(result,2)    
    print("Output>>>>>>>>>>>>>>>>>>>>!!!!!!!!",result)
    return render(request, 'pc/doc_compare.html',{'result': result})
    

#two text compare(.txt, .docx)
def twofilecompare1(request):
    value1 = ''
    value2 = ''
    if (str(request.FILES['docfile1'])).endswith(".txt") and (str(request.FILES['docfile2'])).endswith(".txt"):
        value1 = str(request.FILES['docfile1'].read())
        value2 = str(request.FILES['docfile2'].read())

    elif (str(request.FILES['docfile1'])).endswith(".docx") and (str(request.FILES['docfile2'])).endswith(".docx"):
        document = Document(request.FILES['docfile1'])
        for para in document.paragraphs:
            value1 += para.text
        document = Document(request.FILES['docfile2'])
        for para in document.paragraphs:
            value2 += para.text

    result = fileSimilarity.findFileSimilarity(value1,value2)
    
    print("Output..................!!!!!!!!",result)
    return render(request, 'pc/doc_compare.html',{'result': result})
